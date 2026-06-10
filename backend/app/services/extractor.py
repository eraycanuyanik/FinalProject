"""Belge → düz metin çıkarma.

Desteklenen biçimler: PDF (pdfplumber), DOCX (python-docx), TXT.
Taranmış (görüntü tabanlı) PDF'lerde metin çıkmazsa OCR'a (tesseract, Türkçe)
düşülür. OCR bağımlılıkları yoksa zarifçe atlanır.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument

# Bu eşiğin altında metin çıkan PDF'i "taranmış" sayıp OCR deneriz.
_OCR_MIN_CHARS = 100

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png", ".webp", ".tiff", ".bmp"}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".tiff", ".bmp"}


class UnsupportedFileType(Exception):
    pass


@dataclass
class ExtractionResult:
    text: str
    method: str          # "pdf", "docx", "txt", "pdf+ocr"
    pages: int           # PDF için sayfa sayısı, diğerlerinde 0
    ocr_used: bool


def _ext(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot != -1 else ""


def _ocr_pdf(content: bytes) -> str:
    """Taranmış PDF için OCR. Bağımlılık yoksa boş döner."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except Exception:  # noqa: BLE001
        return ""

    pages = convert_from_bytes(content, dpi=200)
    parts: list[str] = []
    for img in pages:
        parts.append(pytesseract.image_to_string(img, lang="tur"))
    return "\n\n".join(parts).strip()


def _extract_pdf(content: bytes) -> ExtractionResult:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    text = "\n\n".join(parts).strip()

    # Metin çıkmadıysa (taranmış belge) OCR dene.
    if len(text) < _OCR_MIN_CHARS:
        ocr_text = _ocr_pdf(content)
        if len(ocr_text) > len(text):
            return ExtractionResult(
                text=ocr_text, method="pdf+ocr", pages=page_count, ocr_used=True
            )

    return ExtractionResult(text=text, method="pdf", pages=page_count, ocr_used=False)


def _extract_docx(content: bytes) -> ExtractionResult:
    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    # Tablolardaki metni de topla.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return ExtractionResult(text=text, method="docx", pages=0, ocr_used=False)


def _extract_image(content: bytes) -> ExtractionResult:
    """Telefonla çekilmiş / taranmış görsel sözleşmeden OCR ile metin çıkarır."""
    import io as _io

    import pytesseract
    from PIL import Image

    img = Image.open(_io.BytesIO(content))
    text = pytesseract.image_to_string(img, lang="tur").strip()
    return ExtractionResult(text=text, method="image+ocr", pages=1, ocr_used=True)


def _extract_txt(content: bytes) -> ExtractionResult:
    for enc in ("utf-8", "utf-8-sig", "iso-8859-9", "cp1254", "latin-1"):
        try:
            text = content.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = content.decode("utf-8", errors="replace")
    return ExtractionResult(text=text.strip(), method="txt", pages=0, ocr_used=False)


def extract_text(filename: str, content: bytes) -> ExtractionResult:
    """Dosya adına göre uygun çıkarıcıyı seçer."""
    ext = _ext(filename)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    if ext == ".txt":
        return _extract_txt(content)
    if ext in _IMAGE_EXTENSIONS:
        return _extract_image(content)
    raise UnsupportedFileType(
        f"Desteklenmeyen dosya türü: '{ext or filename}'. "
        f"Desteklenenler: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )
